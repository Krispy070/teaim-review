import io, datetime as dt, re
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from .supabase_client import get_supabase_client
from .db import get_conn
import docx

router = APIRouter()

# simple helpers
def _txt(doc_bytes: bytes) -> str:
    d = docx.Document(io.BytesIO(doc_bytes))
    text_parts = []
    
    # Extract paragraph text
    for p in d.paragraphs:
        if p.text.strip():
            text_parts.append(p.text.strip())
    
    # Extract table text (many SOWs have critical data in tables)
    for table in d.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)

def _has(line: str, *keys): 
    s=line.lower(); return all(k.lower() in s for k in keys)

def parse_sow(text: str):
    out = { "phases": [], "workstreams": [], "integrations": [], "roles": [] }

    # phases (look for "Plan Stage", "Architect Stage", etc.)
    for k in ["Plan Stage","Architect Stage","Configure and Prototype Stage","Test Stage","Deploy Stage"]:
        if k in text: out["phases"].append({"name": k.split(" Stage")[0], "order": len(out["phases"])})
    
    # functional areas table (line-based heuristic)
    ws_set = set()
    for line in text.splitlines():
        if _has(line,"HCM") or _has(line,"Core HCM"): ws_set.add("HCM")
        if _has(line,"Payroll"): ws_set.add("Payroll")
        if _has(line,"Benefits"): ws_set.add("Benefits")
        if _has(line,"Absence"): ws_set.add("Time & Absence")
        if _has(line,"Recruiting"): ws_set.add("Recruiting")
        if _has(line,"Talent"): ws_set.add("Talent")
        if _has(line,"Advanced Compensation"): ws_set.add("Advanced Compensation")
        if _has(line,"Time Tracking"): ws_set.add("Time Tracking")
        if _has(line,"Core Financials") or _has(line,"Financial Accounting"): ws_set.add("Finance")
        if _has(line,"Grants"): ws_set.add("Grants")
        if _has(line,"Procurement") or _has(line,"Spend Management"): ws_set.add("Procurement")
        if _has(line,"Supplier Accounts"): ws_set.add("Supplier Accounts")
        if _has(line,"Expenses"): ws_set.add("Expenses")
        if _has(line,"Reporting"): ws_set.add("Reporting/Prism")
        if _has(line,"Security"): ws_set.add("Security")
        if _has(line,"Data Conversion"): ws_set.add("Data Conversion")
        if _has(line,"Integrations"): ws_set.add("Integrations")
        if _has(line,"Cutover") or _has(line,"Go Live"): ws_set.add("Cutover")
    out["workstreams"] = [{"name": n, "description": ""} for n in sorted(ws_set)]

    # integrations (scan appendix A table lines)
    for line in text.splitlines():
        if "→" in line or "-" in line or "Cloud Connect" in line or "Integration" in line:
            if any(x in line for x in ["834","TIAA","Tax Filing","Bank","PNC","M&T","Directory","Student","COBRA","ACH","BAI2","RaisersEdge"]):
                out["integrations"].append(line.strip())

    # roles (Teams & Resources)
    for role in ["Executive Sponsor","Engagement Manager","Project Manager","Architect","Functional Lead","Testing Manager","Change Manager","Data Conversion Lead","Workday Administrator"]:
        if role in text:
            out["roles"].append({"role": role, "name": "", "email": ""})

    return out

@router.post("/sow/ingest")
async def sow_ingest(org_id: str = Form(...), project_id: str = Form(...), file: UploadFile = File(...)):
    # Validate file type and size
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are supported")
    if file.size and file.size > 50 * 1024 * 1024:  # 50MB limit
        raise HTTPException(status_code=400, detail="File size too large (max 50MB)")
    data = await file.read()
    text = _txt(data)

    payload = parse_sow(text)

    # 1) Store phases as episodic memories
    for phase in payload["phases"]:
        try:
            sb = get_supabase_client()
            sb.table("mem_entries").insert({
                "org_id": org_id, "project_id": project_id, "type": "episodic",
                "title": "sow_phase", "body": f"{phase['name']}|order:{phase['order']}"
            }).execute()
        except Exception:
            from .db import get_conn
            with get_conn() as conn, conn.cursor() as cur:
                cur.execute("""insert into mem_entries (org_id, project_id, type, title, body)
                               values (%s,%s,'episodic','sow_phase',%s)""",
                           (org_id, project_id, f"{phase['name']}|order:{phase['order']}"))

    # 2) create workstreams (up to 30)
    # upsert workstreams
    ws_items = [{"name": w["name"], "description": w.get("description","")} for w in payload["workstreams"]][:30]
    try:
        # prefer REST, then psycopg fallback
        sb = get_supabase_client()
        sb.table("workstreams").update({"is_active": False}).eq("org_id",org_id).eq("project_id",project_id).execute()
        for i,it in enumerate(ws_items):
            sb.table("workstreams").insert({
              "org_id": org_id, "project_id": project_id, "name": it["name"],
              "description": it["description"], "sort_order": i, "is_active": True
            }).execute()
    except Exception:
        from .db import get_conn
        with get_conn() as conn, conn.cursor() as cur:
            cur.execute("update workstreams set is_active=false where org_id=%s and project_id=%s",(org_id,project_id))
            for i,it in enumerate(ws_items):
                cur.execute("""insert into workstreams (org_id,project_id,name,description,sort_order,is_active)
                               values (%s,%s,%s,%s,%s,true)""",
                               (org_id,project_id,it["name"],it["description"],i))

    # 3) seed roles as contacts (blank emails to be filled later)
    for r in payload["roles"]:
        try:
            sb = get_supabase_client()
            sb.table("project_contacts").insert({
              "org_id": org_id, "project_id": project_id, "name": r["role"], "email": f"{r['role'].replace(' ','').lower()}@todo",
              "role": r["role"], "workstream": ""
            }).execute()
        except Exception:
            # Fallback to psycopg
            from .db import get_conn
            with get_conn() as conn, conn.cursor() as cur:
                try:
                    cur.execute("""insert into project_contacts (org_id, project_id, name, email, role, workstream)
                                   values (%s,%s,%s,%s,%s,%s)""",
                               (org_id, project_id, r["role"], f"{r['role'].replace(' ','').lower()}@todo", r["role"], ""))
                except Exception: pass

    # 4) stash integrations in mem entries (semantic) for Integrations page seed
    try:
        from .db import get_conn
        with get_conn() as conn, conn.cursor() as cur:
            for line in payload["integrations"][:50]:
                cur.execute("""insert into mem_entries (org_id, project_id, type, title, body)
                               values (%s,%s,'semantic','integration_in_scope',%s)""",
                              (org_id, project_id, line[:4000]))
    except Exception: pass

    return {"ok": True, "phases": len(payload["phases"]), "workstreams": len(ws_items), "roles": len(payload["roles"]), "integrations": len(payload["integrations"])}